import streamlit as st
import calendar
import pandas as pd
from datetime import datetime
from db import get_connection
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

def set_cell_shading(cell, fill):
    """Ajoute une couleur de fond à une cellule Word"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def generate_word(selected_user, selected_projet, selected_month, selected_year, df, total_jours, mois_fr, logo_path=None):
    doc = Document()

    title_table = doc.add_table(rows=1, cols=2)
    title_table.autofit = False

    if logo_path:
        cell_logo = title_table.cell(0, 0)
        paragraph_logo = cell_logo.paragraphs[0]
        run_logo = paragraph_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1))
    else:
        title_table.cell(0, 0).text = ""

    cell_title = title_table.cell(0, 1)
    paragraph_title = cell_title.paragraphs[0]
    run_title = paragraph_title.add_run("Compte rendu d'activité")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    paragraph_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    header_table = doc.add_table(rows=2, cols=2)
    header_table.style = "Table Grid"
    header_table.autofit = True

    header_table.cell(0, 0).text = f"Consultant : {selected_user}"
    header_table.cell(0, 1).text = "Client : "
    header_table.cell(1, 0).text = f"Mois : {mois_fr[selected_month]} {selected_year}"
    header_table.cell(1, 1).text = f"Projet : {selected_projet}"

    doc.add_paragraph("") 

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["JOUR", "DATE", "PRESENCE", "COMMENTAIRE"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        set_cell_shading(hdr_cells[i], "080686") 

    for _, row in df.iterrows():
        jour = row["Jour"]
        date = row["Date"]
        presence = row["Présence"]

        row_cells = table.add_row().cells

        if jour in ["Samedi", "Dimanche"]:
            row_cells[0].text = ""
            row_cells[1].text = date
            row_cells[2].text = ""
            row_cells[3].text = ""
        else:
            row_cells[0].text = jour
            row_cells[1].text = date
            row_cells[2].text = str(presence)
            row_cells[3].text = "" 

    row_cells = table.add_row().cells
    row_cells[0].text = "Total"
    row_cells[1].text = ""
    row_cells[2].text = str(total_jours)
    row_cells[3].text = ""
    for i in range(4):
        run = row_cells[i].paragraphs[0].runs[0]
        run.font.bold = True

    doc.add_paragraph("\nLe                              à\n")
    pied = doc.add_paragraph("Signature consultant                      Signature client")
    pied.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def show_compte_rendu_activite():
    # Droits d'accès
    if "role" not in st.session_state or st.session_state["role"] != "admin":
        st.error("Accès réservé à l'administrateur.")
        st.stop()

    st.markdown("""
        <style>
            .stApp { margin-top: 0; padding: 0; background-color: #f2f2f2; }
            header { visibility: hidden; height: 0px; }
            .stButton>button, .stDownloadButton>button {
                background-color: #080686 !important;
                color: white !important;
                border-radius: 5px;
                font-size: 16px;
            }
            .stButton>button:hover, .stDownloadButton>button:hover {
                background-color: white !important;
                color: #080686 !important;
                border: 1px solid #080686 !important;
            }
            #title {
                position: fixed; top: 0; left: 200px;
                background-color: #ffffff; z-index: 999;
                width:100%; padding: 10px 20px; margin-bottom: 0px;
                border-bottom: 2px solid #000060;
            }
            #title h4 {
                font-size: 29px; font-weight: 600;
                color: #000060; margin-left: 730px;
                font-family: 'New Icon'; padding-bottom: 6px;
            }
            table {
                width: 100%;
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                font-size: 14px; color: #222;
            }
            th {
                background-color: #080686; color: white;
                font-weight: 700; padding: 10px 12px;
                text-align: center !important; user-select: none;
            }
            td {
                background-color: white;
                padding: 10px 12px;
                text-align: center;
                vertical-align: middle;
            }
            tbody tr:nth-child(odd) td { background-color: #f9f9fb; }
            tbody tr:hover td { background-color: #e6f0ff; cursor: pointer; }
        </style>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div id="title">
            <h4> Compte rendu d'activité </h4>  
        </div>
    """, unsafe_allow_html=True)

    # Connexion DB
    conn = get_connection()
    cursor = conn.cursor()

    # Charger les utilisateurs
    cursor.execute("SELECT id, name FROM users")
    users = cursor.fetchall()
    user_dict = {name: uid for uid, name in users}

    # Filtres
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        selected_user = st.selectbox("🧑‍💼 Consultant", list(user_dict.keys()))

    with col2:
        cursor.execute("""
            SELECT p.id, p.nom
            FROM projets p
            JOIN attribution_projet ap ON ap.projet_id = p.id
            WHERE ap.user_id = %s
        """, (user_dict[selected_user],))
        projets = cursor.fetchall()
        projet_dict = {nom: pid for pid, nom in projets}
        selected_projet = st.selectbox("💼 Projet", list(projet_dict.keys()) if projets else [])

    with col3:
        mois_fr = {
            1: "Janvier", 2: "Février", 3: "Mars", 4: "Avril",
            5: "Mai", 6: "Juin", 7: "Juillet", 8: "Août",
            9: "Septembre", 10: "Octobre", 11: "Novembre", 12: "Décembre"
        }
        current_month = datetime.now().month
        selected_month = st.selectbox(
            "📅 Mois",
            list(mois_fr.keys()),
            format_func=lambda x: mois_fr[x],
            index=current_month - 1
        )

    with col4:
        years = list(range(2020, 2031))
        current_year = datetime.now().year
        selected_year = st.selectbox("🗓️ Année", years, index=years.index(current_year))

    if selected_user and selected_projet:
        nb_days = calendar.monthrange(selected_year, selected_month)[1]
        days = [datetime(selected_year, selected_month, d) for d in range(1, nb_days + 1)]

        cursor.execute("""
            SELECT date, valeur 
            FROM feuille_temps
            WHERE user_id = %s
              AND EXTRACT(YEAR FROM date) = %s
              AND EXTRACT(MONTH FROM date) = %s
        """, (user_dict[selected_user], selected_year, selected_month))
        ft_data = {row[0].strftime("%Y-%m-%d"): row[1] for row in cursor.fetchall()}

        rows = []
        jours_fr = {
            "Monday": "Lundi", "Tuesday": "Mardi", "Wednesday": "Mercredi",
            "Thursday": "Jeudi", "Friday": "Vendredi", "Saturday": "Samedi", "Sunday": "Dimanche"
        }
        for d in days:
            jour = jours_fr[d.strftime("%A")]
            date_str = d.strftime("%Y-%m-%d")
            presence = ft_data.get(date_str, 0)
            date_affiche = f"{d.day} {mois_fr[selected_month]} {d.year}"
            rows.append([jour, date_affiche, presence])

        df = pd.DataFrame(rows, columns=["Jour", "Date", "Présence"])
        total_jours = df["Présence"].sum()

        titre_tableau = f"📑 Compte rendu d'activité de {selected_user} pour {selected_projet} - {mois_fr[selected_month]} {selected_year}"

        st.markdown(f"""
            <h3 style='
                color: #006633;
                font-family: "Segoe UI";
                padding-top: 20px;
                padding-bottom: 20px;
                text-align: center;
            '>{titre_tableau}</h3>
        """, unsafe_allow_html=True)

        table_html = df.to_html(index=False, classes="styled-table", escape=False)

        st.markdown(table_html, unsafe_allow_html=True)

        st.markdown(f"""
            <h5 style='
                color: #006633;
                font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif;
                font-weight: 600;
                padding-top: 10px;
            '>Nombre total de jours : {total_jours}</h5>
        """, unsafe_allow_html=True)

        st.download_button(
            label="Générer le compte rendu d'activité",
            data=generate_word(
                selected_user,
                selected_projet,
                selected_month,
                selected_year,
                df,
                total_jours,
                mois_fr,
                logo_path="assets/logo.png"
            ),
            file_name=f"Compte_rendu_{selected_user}_{selected_projet}_{mois_fr[selected_month]}_{selected_year}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    cursor.close()
    conn.close()
